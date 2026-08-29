"""Convert the manuscript to a readable .docx for collaborators who do not use LaTeX.

This is a reading copy, not a submission artifact. Pandoc is not installed here, so the
conversion is done directly: sections and subsections become Word headings, paragraphs become
paragraphs, tables become Word tables, and LaTeX markup is unwrapped rather than dropped so
that emphasis and inline math survive as legible text.

What it deliberately does not do is pretend to be the typeset paper. Citation keys render as
bracketed keys rather than numbered references, the TikZ figure is described rather than drawn,
and equations appear as their source. The PDF remains the faithful rendering, and the header of
the generated document says so.

    python tex_to_docx.py <main.tex> <out.docx>
"""
from __future__ import annotations

import io
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def strip_markup(t: str) -> str:
    """Unwrap LaTeX into readable text, preserving content rather than deleting it."""
    t = re.sub(r"(?<!\\)%.*", "", t)                       # comments, not escaped percents
    t = t.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_").replace(r"\#", "#")
    t = re.sub(r"\\(?:emph|textbf|textit|texttt|textsc)\{([^{}]*)\}", r"\1", t)
    t = re.sub(r"\\(?:cite|citep|citet)\{([^}]*)\}", r"[\1]", t)
    t = re.sub(r"\\(?:ref|autoref|eqref)\{([^}]*)\}", r"(\1)", t)
    t = re.sub(r"\\label\{[^}]*\}", "", t)
    t = re.sub(r"\\(?:IEEEPARstart)\{([^}]*)\}\{([^}]*)\}", r"\1\2", t)
    t = t.replace("---", "\u2014").replace("--", "\u2013")
    t = t.replace(r"\ldots", "...").replace(r"\dag", "\u2020").replace(r"\ddag", "\u2021")
    t = t.replace(r"\le", "\u2264").replace(r"\ge", "\u2265").replace(r"\times", "\u00d7")
    t = t.replace(r"\alpha", "\u03b1").replace(r"\Delta", "\u0394").replace(r"\pi", "\u03c0")
    t = t.replace(r"\lesssim", "\u2272").replace(r"\equiv", "\u2261").replace(r"\in", "\u2208")
    # Table glyphs. These carry the comparison in the related-work table, so letting the
    # generic control-sequence sweep delete them would empty that table of its content.
    t = t.replace(r"\checkmark", "\u2713").replace(r"\textasciitilde", "~")
    t = t.replace(r"\ding{51}", "\u2713").replace(r"\ding{55}", "\u2717")
    t = t.replace(r"\Diamond", "\u25c7").replace(r"\diamond", "\u25c7")
    t = re.sub(r"\$([^$]*)\$", r"\1", t)                   # inline math -> its content
    t = t.replace("{=}", "=").replace("{:}", ":")
    t = re.sub(r"\\[a-zA-Z]+\*?", "", t)                   # any remaining control sequence
    t = t.replace("{", "").replace("}", "").replace("~", " ")
    return " ".join(t.split())


def parse_tables(body: str) -> dict:
    """Map a table's caption text to its rows, so tables survive as tables."""
    out = {}
    for m in re.finditer(r"\\begin\{table\*?\}(.*?)\\end\{table\*?\}", body, re.S):
        block = m.group(1)
        cap = re.search(r"\\caption\{(.*?)\}\s*\n", block, re.S)
        caption = strip_markup(cap.group(1)) if cap else "Table"
        tab = re.search(r"\\begin\{tabular\}\{[^}]*\}(.*?)\\end\{tabular\}", block, re.S)
        if not tab:
            continue
        rows = []
        for line in tab.group(1).split(r"\\"):
            line = re.sub(r"\\(?:top|mid|bottom)rule", "", line)
            line = re.sub(r"\\multicolumn\{\d+\}\{[^}]*\}\{([^}]*)\}", r"\1", line)
            cells = [strip_markup(c) for c in line.split("&")]
            if any(c for c in cells):
                rows.append(cells)
        if rows:
            out[m.start()] = (caption, rows)
    return out


def main(tex_path: str, out_path: str) -> int:
    src = io.open(tex_path, encoding="utf-8").read()

    title = strip_markup(re.search(r"\\title\{(.*?)\}\s*\n", src, re.S).group(1))
    abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", src, re.S).group(1)
    keywords = re.search(r"\\begin\{IEEEkeywords\}(.*?)\\end\{IEEEkeywords\}", src, re.S)

    start = src.index(r"\section{Introduction}")
    end = src.index(r"\bibliographystyle")
    body = src[start:end]
    tables = parse_tables(body)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(title, level=0)

    note = doc.add_paragraph()
    note.add_run(
        "Reading copy generated from main.tex. The typeset PDF is the faithful rendering: "
        "here, citations appear as bracketed keys, equations as their source, and the pipeline "
        "figure is described rather than drawn."
    ).italic = True

    doc.add_heading("Abstract", level=1)
    for para in [p for p in abstract.split("\n\n") if p.strip()]:
        doc.add_paragraph(strip_markup(para))
    if keywords:
        p = doc.add_paragraph()
        p.add_run("Index terms\u2014").bold = True
        p.add_run(strip_markup(keywords.group(1)))

    # Walk the body, emitting headings, tables and paragraphs in source order.
    chunks = re.split(r"(\\(?:sub)*section\*?\{[^}]*\})", body)
    emitted_tables = set()
    for chunk in chunks:
        m = re.match(r"\\(sub)*section\*?\{([^}]*)\}", chunk)
        if m:
            level = 1 + len(re.findall("sub", m.group(0)))
            doc.add_heading(strip_markup(m.group(2)), level=min(level, 4))
            continue
        for off, (caption, rows) in sorted(tables.items()):
            if off in emitted_tables or chunk.find(caption[:28]) < 0:
                continue
            emitted_tables.add(off)
            t = doc.add_table(rows=0, cols=max(len(r) for r in rows))
            t.style = "Light Grid Accent 1"
            for r in rows:
                cells = t.add_row().cells
                for i, val in enumerate(r[:len(cells)]):
                    cells[i].text = val
            cp = doc.add_paragraph()
            cp.add_run(caption).italic = True
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean = re.sub(r"\\begin\{table\*?\}.*?\\end\{table\*?\}", "", chunk, flags=re.S)
        clean = re.sub(r"\\begin\{figure\*?\}.*?\\end\{figure\*?\}",
                       "[Figure 1: the six-stage evaluation trajectory. Five model-facing "
                       "decisions surround one deterministic environment step; see the PDF.]",
                       clean, flags=re.S)
        clean = re.sub(r"\\begin\{(itemize|enumerate)\}|\\end\{(itemize|enumerate)\}", "", clean)
        for para in [p for p in clean.split("\n\n") if p.strip()]:
            items = [i for i in re.split(r"\\item\s", para) if i.strip()]
            if len(items) > 1 or para.strip().startswith(r"\item"):
                for it in items:
                    txt = strip_markup(it)
                    if txt:
                        doc.add_paragraph(txt, style="List Bullet")
            else:
                txt = strip_markup(para)
                if txt:
                    doc.add_paragraph(txt)

    doc.save(out_path)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print("wrote %s" % os.path.basename(out_path))
    print("  paragraphs %d, tables %d, words %d"
          % (len(doc.paragraphs), len(doc.tables), words))
    return 0


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(2)
    sys.exit(main(sys.argv[1], sys.argv[2]))
